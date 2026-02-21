import streamlit as st
import requests
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# 1. إعداد الصفحة والتصميم الملكي (أسود وذهبي)
st.set_page_config(page_title="مساعد الأستاذة نور", page_icon="📐", layout="wide")

# --- ضعي مفتاح Groq الخاص بكِ هنا ---
GROQ_API_KEY = "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" 
# -----------------------------------

# 2. تصميم CSS الاحترافي (الهيكلية الجميلة)
st.markdown("""
    <style>
    [data-testid="stSidebar"] { background-color: #1e2130; border-right: 2px solid gold; }
    .stButton>button {
        width: 100%; border-radius: 8px; height: 3.5em;
        background-color: #3498db; color: white; font-weight: bold; border: none;
    }
    .stButton>button:hover { background-color: #2980b9; border: 1px solid white; }
    .main-header {
        background-color: #1e1e1e; color: #FFD700; text-align: center;
        padding: 25px; border-radius: 15px; border-bottom: 4px solid #FFD700;
        margin-bottom: 20px; font-family: 'Arial';
    }
    .reportview-container { direction: rtl; }
    div.stMarkdown { text-align: right; }
    </style>
    """, unsafe_allow_html=True)

# 3. دالة إنشاء ملف Word يدعم العربية 100%
def create_word_rtl(text, topic):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(14)
    
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = heading.add_run(f"خطة درس: {topic}")
    run.bold = True
    run.font.size = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(text)
    
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# 4. القائمة الجانبية (Sidebar) مع الشعار
with st.sidebar:
    # عرض شعار المدرسة
    st.image("https://raw.githubusercontent.com/NoorMath/NoorApp.py/main/school_logo.jpg", width=160)
    st.markdown("<h2 style='color: gold; text-align: center;'>لوحة التحكم</h2>", unsafe_allow_html=True)
    st.write("---")
    grade = st.selectbox("اختر المرحلة الدراسية:", ["الثالث المتوسط", "الرابع العلمي", "الخامس العلمي"])
    st.write("")
    if st.button("🔄 تحديث النظام"):
        st.rerun()
    st.write("---")
    st.markdown("<p style='text-align: center; color: #ccc;'>إعداد الأستاذة نور محمد حسن</p>", unsafe_allow_html=True)

# 5. الواجهة الرئيسية
st.markdown("<div class='main-header'><h1>نظام تحضير دروس الرياضيات الذكي</h1><p>ثانوية خير الأنام للبنين</p></div>", unsafe_allow_html=True)

topic = st.text_input("📝 أدخلي موضوع الدرس المراد تحضيره:", placeholder="مثلاً: حل المعادلات التربيعية بالمربع الكامل")

# 6. توليد الخطة
if st.button("توليد الخطة النموذجية 🚀"):
    if topic:
        with st.spinner("⏳ جاري صياغة الخطة باللغة العربية الفصحى..."):
            try:
                url = "https://api.groq.com/openai/v1/chat/completions"
                headers = {"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"}
                data = {
                    "model": "llama-3.3-70b-versatile",
                    "messages": [
                        {"role": "system", "content": "أنت خبير تربوي عراقي. اكتب الخطة باللغة العربية فقط وبشكل منظم جداً (أهداف، وسائل، عرض، تقويم)."},
                        {"role": "user", "content": f"اكتب خطة درس رياضيات لصف {grade} عن {topic}. وفي النهاية اكتب: إعداد الأستاذة نور محمد حسن."}
                    ]
                }
                response = requests.post(url, headers=headers, json=data)
                plan = response.json()['choices'][0]['message']['content']
                st.session_state['current_plan'] = plan
                st.session_state['topic_name'] = topic
            except Exception as e:
                st.error("تأكدي من مفتاح Groq")

# 7. عرض النتائج وأزرار التحميل
if 'current_plan' in st.session_state:
    st.markdown("### 📄 الخطة الدراسية المستخرجة:")
    st.info(st.session_state['current_plan'])
    
    st.write("---")
    st.subheader("📥 تحميل الخطة بصيغة واضحة:")
    
    c1, c2 = st.columns(2)
    with c1:
        # تحميل وورد (يدعم التنسيق العربي)
        st.download_button(
            label="📥 تحميل ملف Word (للطباعة)",
            data=create_word_rtl(st.session_state['current_plan'], st.session_state['topic_name']),
            file_name=f"خطة_{st.session_state['topic_name']}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    with c2:
        # تحميل نص (سريع وواضح جداً للموبايل)
        st.download_button(
            label="📝 تحميل ملف نصي (Text)",
            data=st.session_state['current_plan'],
            file_name=f"خطة_{st.session_state['topic_name']}.txt",
            mime="text/plain"
        )

st.markdown("<br><hr><center><b>تم التطوير بواسطة الأستاذة نور محمد حسن © 2026</b></center>", unsafe_allow_html=True)
